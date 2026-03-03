#!/usr/bin/env python3
from __future__ import annotations

import json
import os
import sqlite3
import sys
import time
import unittest
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
except Exception:  # pragma: no cover - dependency may be missing in minimal envs
    Document = None  # type: ignore[assignment]

try:
    from dotenv import load_dotenv
except Exception:  # pragma: no cover - optional convenience only
    load_dotenv = None


class Colors:
    RESET = "\033[0m"
    BOLD = "\033[1m"
    RED = "\033[31m"
    GREEN = "\033[32m"
    YELLOW = "\033[33m"
    CYAN = "\033[36m"


@dataclass
class RunResult:
    fmt: str
    input_file: str
    output_file: str
    elapsed_seconds: float
    terms_found: int
    terms_total: int
    hit_rate: float
    success: bool
    error: str | None = None


class TestRunner:
    def __init__(
        self,
        project_root: Path,
        fixture_dir: Path,
        result_dir: Path,
        db_path: Path,
        src_lang: str = "cn",
        tgt_lang: str = "en",
        engine: str = "xfyun",
        domain: str | None = None,
        max_workers: int = 4,
    ):
        self.project_root = project_root
        self.fixture_dir = fixture_dir
        self.result_dir = result_dir
        self.db_path = db_path
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.engine = engine
        self.domain = domain
        self.max_workers = max(1, max_workers)

        self._results_cache: dict[str, RunResult] = {}
        self._term_pairs = self._load_term_pairs()

        sys.path.insert(0, str(self.project_root / "src"))
        from file_translator.core import translate_file

        self.translate_file = translate_file

    def _load_term_pairs(self) -> list[tuple[str, str]]:
        if not self.db_path.exists():
            raise RuntimeError(f"Terminology DB not found: {self.db_path}")

        conn = sqlite3.connect(self.db_path)
        try:
            src = self.src_lang.strip().lower()
            tgt = self.tgt_lang.strip().lower()
            rows = conn.execute(
                """
                SELECT DISTINCT ls.text AS src_text, lt.text AS tgt_text
                FROM term_concept c
                JOIN term_lexeme ls ON ls.concept_id = c.id
                JOIN term_lexeme lt ON lt.concept_id = c.id
                WHERE ls.lang = ? AND lt.lang = ?
                  AND ls.status = 'approved' AND lt.status = 'approved'
                  AND (? IS NULL OR c.domain = ?)
                """,
                (src, tgt, self.domain, self.domain),
            ).fetchall()
        finally:
            conn.close()

        pairs = []
        for src_text, tgt_text in rows:
            if src_text and tgt_text:
                pairs.append((str(src_text), str(tgt_text)))
        return pairs

    def _load_docx_text(self, path: Path) -> str:
        doc = Document(str(path))
        chunks: list[str] = []

        for p in doc.paragraphs:
            if p.text:
                chunks.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text:
                            chunks.append(p.text)

        return "\n".join(chunks)

    def _load_text_content(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8")
        if suffix == ".json":
            obj = json.loads(path.read_text(encoding="utf-8"))
            return self._flatten_json_strings(obj)
        if suffix == ".docx":
            return self._load_docx_text(path)
        raise RuntimeError(f"Unsupported format: {suffix}")

    def _flatten_json_strings(self, node: Any) -> str:
        out: list[str] = []

        def walk(cur: Any):
            if isinstance(cur, dict):
                for v in cur.values():
                    walk(v)
            elif isinstance(cur, list):
                for v in cur:
                    walk(v)
            elif isinstance(cur, str):
                out.append(cur)

        walk(node)
        return "\n".join(out)

    def _output_path_for(self, in_path: Path) -> Path:
        out_name = f"{in_path.stem}.translated{in_path.suffix.lower()}"
        return self.result_dir / out_name

    def _calc_term_hit_rate(self, source_text: str, translated_text: str) -> tuple[int, int, float]:
        opportunities: list[tuple[str, str]] = []
        for src_term, tgt_term in self._term_pairs:
            if src_term in source_text:
                opportunities.append((src_term, tgt_term))

        total = len(opportunities)
        if total == 0:
            return 0, 0, 0.0

        found = sum(1 for _, tgt_term in opportunities if tgt_term in translated_text)
        rate = found / total * 100.0
        return found, total, rate

    def run_format(self, filename: str) -> RunResult:
        suffix = Path(filename).suffix.lower().lstrip(".")
        if suffix in self._results_cache:
            return self._results_cache[suffix]

        in_path = self.fixture_dir / filename
        out_path = self._output_path_for(in_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

        source_text = self._load_text_content(in_path)

        started = time.perf_counter()
        error: str | None = None
        success = False
        try:
            self.translate_file(
                in_path=in_path,
                out_path=out_path,
                src=self.src_lang,
                tgt=self.tgt_lang,
                engine=self.engine,
                domain=self.domain,
                max_workers=self.max_workers,
            )
            success = out_path.exists()
        except Exception as exc:
            error = str(exc)
            success = False

        elapsed = time.perf_counter() - started

        translated_text = ""
        terms_found = 0
        terms_total = 0
        hit_rate = 0.0

        if success:
            translated_text = self._load_text_content(out_path)
            terms_found, terms_total, hit_rate = self._calc_term_hit_rate(source_text, translated_text)

        result = RunResult(
            fmt=suffix,
            input_file=str(in_path),
            output_file=str(out_path),
            elapsed_seconds=elapsed,
            terms_found=terms_found,
            terms_total=terms_total,
            hit_rate=hit_rate,
            success=success,
            error=error,
        )
        self._results_cache[suffix] = result
        return result

    def run_all(self) -> list[RunResult]:
        order = ["sample.txt", "sample.md", "sample.json", "sample.docx"]
        return [self.run_format(name) for name in order]

    def print_summary(self, results: list[RunResult]) -> None:
        print(f"\n{Colors.BOLD}{Colors.CYAN}=== Multi-Format Translation Summary ==={Colors.RESET}")
        print(f"Engine: {self.engine} | Lang: {self.src_lang}->{self.tgt_lang} | DB: {self.db_path}")
        print("+--------+----------+-------------+----------+--------+")
        print("| Format | Time(s)  | Terms Found | Hit Rate | Status |")
        print("+--------+----------+-------------+----------+--------+")
        for r in results:
            status = f"{Colors.GREEN}PASS{Colors.RESET}" if r.success else f"{Colors.RED}FAIL{Colors.RESET}"
            tf = f"{r.terms_found}/{r.terms_total}"
            print(f"| {r.fmt:<6} | {r.elapsed_seconds:>8.3f} | {tf:>11} | {r.hit_rate:>7.2f}% | {status:>6} |")
        print("+--------+----------+-------------+----------+--------+")

        failed = [r for r in results if not r.success]
        if failed:
            print(f"{Colors.RED}Failures:{Colors.RESET}")
            for r in failed:
                print(f"- {r.fmt}: {r.error or 'Unknown error'}")
        else:
            print(f"{Colors.GREEN}All format tests passed.{Colors.RESET}")

    def save_report(self, results: list[RunResult]) -> tuple[Path, Path]:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.result_dir.mkdir(parents=True, exist_ok=True)

        json_path = self.result_dir / f"summary_{ts}.json"
        txt_path = self.result_dir / f"summary_{ts}.txt"

        payload = {
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "engine": self.engine,
            "src_lang": self.src_lang,
            "tgt_lang": self.tgt_lang,
            "db_path": str(self.db_path),
            "results": [asdict(r) for r in results],
        }
        json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

        lines = [
            "Multi-Format Translation Summary",
            f"Generated: {payload['generated_at']}",
            f"Engine: {self.engine}",
            f"Lang: {self.src_lang}->{self.tgt_lang}",
            "",
            "Format | Time(s) | Terms Found | Hit Rate | Status",
        ]
        for r in results:
            status = "PASS" if r.success else "FAIL"
            lines.append(
                f"{r.fmt} | {r.elapsed_seconds:.3f} | {r.terms_found}/{r.terms_total} | {r.hit_rate:.2f}% | {status}"
            )
            if r.error:
                lines.append(f"  error: {r.error}")

        txt_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        return json_path, txt_path


def create_fixture_files(fixture_dir: Path) -> None:
    fixture_dir.mkdir(parents=True, exist_ok=True)

    sample_txt = """2D/3D切换
HDMI输入
安装工具
本段用于测试txt格式翻译。
"""
    (fixture_dir / "sample.txt").write_text(sample_txt, encoding="utf-8")

    sample_md = """# 手术系统翻译测试

## 功能列表

- 2D/3D切换
- HDMI输出
- SDI输入

请确保术语翻译准确，并保留 Markdown 结构。
"""
    (fixture_dir / "sample.md").write_text(sample_md, encoding="utf-8")

    sample_json = {
        "title": "翻译质量检查",
        "description": "用于测试JSON字符串翻译。",
        "items": ["U盘", "控制软件", "电凝器械"],
        "metadata": {
            "note": "请保留JSON键名不变",
            "owner": "测试团队",
        },
    }
    (fixture_dir / "sample.json").write_text(json.dumps(sample_json, ensure_ascii=False, indent=2), encoding="utf-8")

    if Document is not None:
        doc = Document()
        doc.add_heading("DOCX 翻译验证", level=1)
        doc.add_paragraph("本文件用于测试 Word 文档翻译。")
        doc.add_paragraph("包含术语：机械臂、控制台臂托、电刀。")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "字段"
        table.cell(0, 1).text = "内容"
        table.cell(1, 0).text = "设备"
        table.cell(1, 1).text = "高清内窥镜"
        doc.save(str(fixture_dir / "sample.docx"))


def _require_xfyun_env(engine: str) -> None:
    if engine.lower() != "xfyun":
        return
    required = ["XFYUN_APP_ID", "XFYUN_API_KEY", "XFYUN_API_SECRET"]
    missing = [k for k in required if not os.getenv(k)]
    if missing:
        raise unittest.SkipTest(
            "Missing required env vars for real xfyun translation: " + ", ".join(missing)
        )


class MultiFormatTranslationTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        script_path = Path(__file__).resolve()
        cls.project_root = script_path.parent.parent

        if load_dotenv is not None:
            load_dotenv(cls.project_root / ".env")

        cls.fixture_dir = cls.project_root / "test" / "fixtures"
        cls.result_dir = cls.project_root / "test" / "results"
        cls.db_path = cls.project_root / "src" / "file_translator" / "translator_terms.db"

        create_fixture_files(cls.fixture_dir)
        cls.result_dir.mkdir(parents=True, exist_ok=True)

        cls.engine = os.getenv("TEST_TRANSLATION_ENGINE", "xfyun")
        cls.src_lang = os.getenv("TEST_TRANSLATION_SRC", "cn")
        cls.tgt_lang = os.getenv("TEST_TRANSLATION_TGT", "en")
        cls.domain = os.getenv("TEST_TRANSLATION_DOMAIN") or None
        cls.max_workers = int(os.getenv("TEST_TRANSLATION_MAX_WORKERS", "4"))

        _require_xfyun_env(cls.engine)
        if Document is None:
            raise unittest.SkipTest("python-docx is not installed. Install dependencies to run DOCX tests.")

        cls.runner = TestRunner(
            project_root=cls.project_root,
            fixture_dir=cls.fixture_dir,
            result_dir=cls.result_dir,
            db_path=cls.db_path,
            src_lang=cls.src_lang,
            tgt_lang=cls.tgt_lang,
            engine=cls.engine,
            domain=cls.domain,
            max_workers=cls.max_workers,
        )

    def _assert_success_result(self, result: RunResult) -> None:
        self.assertTrue(result.success, msg=f"{result.fmt} translation failed: {result.error}")
        self.assertTrue(Path(result.output_file).exists(), msg=f"Output file missing for {result.fmt}")
        self.assertGreater(result.elapsed_seconds, 0.0)

    def test_txt_translation(self) -> None:
        result = self.runner.run_format("sample.txt")
        self._assert_success_result(result)

    def test_md_translation(self) -> None:
        result = self.runner.run_format("sample.md")
        self._assert_success_result(result)

    def test_json_translation(self) -> None:
        result = self.runner.run_format("sample.json")
        self._assert_success_result(result)

    def test_docx_translation(self) -> None:
        if Document is None:
            self.skipTest("python-docx is not installed")
        result = self.runner.run_format("sample.docx")
        self._assert_success_result(result)

    def test_all_formats(self) -> None:
        results = self.runner.run_all()
        self.runner.print_summary(results)
        json_report, txt_report = self.runner.save_report(results)
        print(f"{Colors.YELLOW}Saved report:{Colors.RESET} {json_report}")
        print(f"{Colors.YELLOW}Saved report:{Colors.RESET} {txt_report}")

        self.assertEqual(len(results), 4)
        self.assertTrue(all(r.success for r in results), msg="One or more format translations failed")


if __name__ == "__main__":
    unittest.main(verbosity=2)
