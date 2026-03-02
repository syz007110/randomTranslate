from __future__ import annotations

import json
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any

from docx import Document


def _translate_many(texts: list[str], translate_fn, max_workers: int = 4) -> list[str]:
    if not texts:
        return []
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        return list(ex.map(translate_fn, texts))


def translate_txt(content: str, translate_fn, max_workers: int = 4) -> str:
    lines = content.splitlines(keepends=True)
    idxs, texts = [], []
    for i, ln in enumerate(lines):
        raw = ln.rstrip("\r\n")
        if raw.strip():
            idxs.append(i)
            texts.append(raw)

    translated = _translate_many(texts, translate_fn, max_workers=max_workers)
    for i, t in zip(idxs, translated):
        ln = lines[i]
        lines[i] = t + ("\n" if ln.endswith("\n") else "")
    return "".join(lines)


def translate_md(content: str, translate_fn, max_workers: int = 4) -> str:
    lines = content.splitlines(keepends=True)
    in_code = False
    idxs, texts = [], []

    for i, ln in enumerate(lines):
        striped = ln.lstrip()
        if striped.startswith("```"):
            in_code = not in_code
            continue
        if in_code or not ln.strip():
            continue
        idxs.append(i)
        texts.append(ln.rstrip("\r\n"))

    translated = _translate_many(texts, translate_fn, max_workers=max_workers)
    for i, t in zip(idxs, translated):
        ln = lines[i]
        lines[i] = t + ("\n" if ln.endswith("\n") else "")
    return "".join(lines)


def _collect_json_strings(node: Any, path=()):
    items = []
    if isinstance(node, dict):
        for k, v in node.items():
            items.extend(_collect_json_strings(v, path + (k,)))
    elif isinstance(node, list):
        for i, v in enumerate(node):
            items.extend(_collect_json_strings(v, path + (i,)))
    elif isinstance(node, str):
        items.append((path, node))
    return items


def _set_json_path(root: Any, path: tuple, value: str):
    node = root
    for p in path[:-1]:
        node = node[p]
    node[path[-1]] = value


def translate_json(content: str, translate_fn, max_workers: int = 4) -> str:
    obj = json.loads(content)
    pairs = _collect_json_strings(obj)
    paths = [p for p, _ in pairs]
    texts = [t for _, t in pairs]
    translated = _translate_many(texts, translate_fn, max_workers=max_workers)
    for p, t in zip(paths, translated):
        _set_json_path(obj, p, t)
    return json.dumps(obj, ensure_ascii=False, indent=2)


def _collect_table_runs(table):
    runs = []
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.text and run.text.strip():
                        runs.append(run)
            # nested tables
            for nested in cell.tables:
                runs.extend(_collect_table_runs(nested))
    return runs


def _collect_docx_runs(doc):
    runs = []
    # body paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text and run.text.strip():
                runs.append(run)

    # body tables
    for table in doc.tables:
        runs.extend(_collect_table_runs(table))

    # headers/footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            for run in p.runs:
                if run.text and run.text.strip():
                    runs.append(run)
        for t in section.header.tables:
            runs.extend(_collect_table_runs(t))
        for p in section.footer.paragraphs:
            for run in p.runs:
                if run.text and run.text.strip():
                    runs.append(run)
        for t in section.footer.tables:
            runs.extend(_collect_table_runs(t))

    return runs


def translate_docx(in_path: Path, out_path: Path, translate_fn, max_workers: int = 4) -> None:
    doc = Document(str(in_path))
    runs = _collect_docx_runs(doc)
    texts = [r.text for r in runs]
    translated = _translate_many(texts, translate_fn, max_workers=max_workers)
    for run, t in zip(runs, translated):
        run.text = t
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
